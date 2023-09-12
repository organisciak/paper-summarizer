import PyPDF2
from tqdm.auto import tqdm
import openai
import re
from yaml_sync import YamlCache
from pathlib import Path
import hashlib
from .utils import DocSplitter, word_diff
import numpy as np
import time
import warnings
from IPython.display import display, Markdown, HTML, clear_output


class PaperSummarizer:
    def __init__(self, text, cache=True, cache_location=None,
                 model='gpt-3.5-turbo', doctype='research paper',
                 drop_refs=True):
        self.cache_location = cache_location
        self.model = model
        self.text = text
        self.splitter = None
        self.chunks = None
        self.doctype = doctype
        self.drop_refs = drop_refs

        assert not (cache and cache_location is None), "Set a location for yaml file with cache_location if cache is True"

        self.splitter, self.chunks = self._split_doc(model)

        if cache:
            self.cache = YamlCache(self.cache_location, mode='rw',
                                   order=['model', 'full_summary', 'questions',
                                          'full_points', 'chunk_outlines'])
        else:
            self.cache = {}
        self.cache['model'] = model

    def question(self, question, model='default', temperature=0, force=False,
                 md=False, stream=False):
        ''' Once the full summary and keypoints are created, you can ask
        ad-hoc questions.
        '''

        summary = self.full_summary()
        pts = self.full_points()

        big_summary = f'''The following is a summary of a {self.doctype} and it's key points.

# SUMMARY
{summary}

# KEY POINTS
{pts}

-------------

Are you ready to answer questions about the {self.doctype}?
'''
        result = self._question(big_summary, question, 'questions',
                                model=model, temperature=temperature,
                                force=force, stream=stream, md=md)

        if md:
            return Markdown(result.replace('\n', '\n\n'))
        else:
            return result

    def _split_doc(self, model, chunksize='auto'):
        '''
        Initialize token splitter and split document into chunks.
        If the model changes, we need to re-split the text - currently this is done temporarily, without changing the class state
        '''
        if model == 'default':
            return self.splitter, self.chunks

        update_chunks = False
        if chunksize == 'auto':
            if '32k' in model:
                chunksize = 2 * 32768 // 3
            elif '16k' in model:
                chunksize = 2 * 16384 // 3
            elif model.startswith('gpt-4'):
                chunksize = 2 * 8192 // 3
            elif model.startswith('gpt-3.5'):
                chunksize = 2 * 4097 // 3
            else:
                print(f'unrecognized model {model}, assuming 4k context')
                chunksize = 2 * 4097 // 3
        else:
            # saved chunks may not match the target chunk size
            update_chunks = True

        if (self.chunks is None) or (self.splitter is None) or (self.splitter.model != model) or update_chunks:
            splitter = DocSplitter(self.text, encoding_model=model, drop_refs=self.drop_refs)
            chunks = splitter.doc_to_chunks(max_size=chunksize)
            return splitter, chunks
        else:
            return self.splitter, self.chunks

    def _stream_response(self, response, draw_every=0, md=False,
                         clear_on_finish=True):
        ''' Print a streaming response, as it comes in.

        Parameters:
        response: the streaming response object from OpenAI
        draw_every: how often to print the response, in seconds. If 0, draw every event.
        md: if True, format the response as Markdown
        clear_on_finish: if True, clear the output when the response is finished.
        '''
        collected_chunks = []
        collected_content = ""
        last_print = time.time()
        role = None

        for chunk in response:
            collected_chunks.append(chunk)
            chunk_message = chunk['choices'][0]['delta']
            collected_content += chunk_message.get('content', '')

            if 'role' in chunk_message:
                role = chunk_message['role']

            if (time.time() - last_print) > draw_every:
                if md:
                    clear_output(wait=True)
                    display(Markdown(collected_content.replace('\n', '\n\n')))
                else:
                    print(chunk_message.get('content', ''), end='')
                last_print = time.time()

        if clear_on_finish:
            clear_output(wait=True)

        return {'role': role, 'content': collected_content}

    def direct_question(self, question, model='default', temperature=0,
                        force=False, md=False, target_chunk=0, target_text=None, no_cache=False,
                        stream=False):
        ''' Ask a question directly of the full text (only the first chunk,
        unless target_chunk='all'). You can also supply text directly, with target_text.

        Usually you want `question`, unless you have a shorter text.
        '''
        if (('direct_questions' in self.cache) and (question in self.cache['direct_questions']) and not force):
            cached_response = self.cache['direct_questions'][question]
            cached_response = "\n".join(cached_response)
            if md:
                cached_response = Markdown(cached_response.replace('\n', '\n\n'))
            return cached_response

        if model == 'default':
            model = self.model
        _, chunks = self._split_doc(model)

        if target_text is not None:
            if target_chunk is not None:
                print("Prioritizing target_text over target_chunk")
            target_chunks = [target_text]
        elif target_chunk == 'all':
            target_chunks = chunks
        else:
            target_chunks = [chunks[min(target_chunk, len(chunks) - 1)]]

        if len(target_chunks) > 1:
            pbar = tqdm(total=len(target_chunks), desc="Collecting Line Item Suggestions", position=1, leave=False)

        direct_q_template = '''The following is a document to talk about.

{}
------------
{}

-------------

Are you ready to answer questions about the document?
'''
        results = []

        for target_text in target_chunks:
            prompt = direct_q_template.format(self.doctype.upper(), target_text)
            result = self._question(prompt, question, 'direct_questions',
                                    model=model,
                                    temperature=temperature,
                                    stream=stream,
                                    md=md,
                                    force=force, no_cache=True)
            results.append(result)
            if len(target_chunks) > 1:
                pbar.update(1)

        final_result = "\n------------\n".join(results)

        if not no_cache:
            if 'direct_questions' not in self.cache:
                self.cache['direct_questions'] = {}
            qs = self.cache['direct_questions']
            qs[question] = final_result.split('\n')
            self.cache['direct_questions'] = qs

        if md:
            return Markdown(final_result.replace('\n', '\n\n'))
        else:
            return final_result

    def _question(self, summary_prompt, question, key, model='default', temperature=0, force=False,
                  no_cache=False, stream=False, md=False, retries=3, cooldown=5):
        if model == 'default':
            model = self.model

        if (key in self.cache) and (question in self.cache[key]) and not force:
            cached_response = self.cache[key][question]
            cached_response = "\n".join(cached_response)
            return cached_response

        big_summary = {"role": "user", "content": summary_prompt}
        bot_affirm = {"role": "assistant", "content": "Yes, I am ready to answer questions about the {self.doctype}."}

        for attempt in range(retries):
            try:
                result = openai.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are a helpful document reader and assistant."},
                        big_summary, bot_affirm,
                        {"role": "user", "content": question}
                    ],
                    stream=stream,
                    temperature=temperature
                )

                if stream:
                    result_msg = self._stream_response(result, md=md)
                else:
                    result_msg = result.choices[0].message

                if not no_cache:
                    if key not in self.cache:
                        self.cache[key] = {}
                    qs = self.cache[key]
                    qs[question] = result_msg.get('content', '').split('\n')
                    self.cache[key] = qs
                    return "\n".join(qs[question])
                else:
                    return result_msg.get('content', '')
            except (openai.error.APIError, openai.error.TryAgain, openai.error.Timeout) as e:
                if attempt < retries - 1:
                    warnings.warn(f"Attempt {attempt + 1} failed with error: {str(e)}. Retrying in {cooldown} seconds...")
                    time.sleep(cooldown)
                    continue
                else:
                    raise e

    def summarize(self, md=False, model='default', protocol=0, force=False):
        ''' Print full summaries'''
        summary = self.full_summary(protocol=protocol, force=force, model=model)
        pts = self.full_points(protocol=protocol, model=model)

        out = f'## Summary\n{summary}\n## Key Points\n{pts}'
        if md:
            display(Markdown(out.replace('\n', '\n\n')))
        else:
            print(out)

    def outline_all_chunks(self, model='default', force=False, protocol=0, debug=False):
        ''' Iterate through chunks and get summaries and key points for each one.

        Protocol is an integer referring to a specific prompt style, and is passed down to outline_chunk.

        '''
        if model == 'default':
            model = self.model
        _, chunks = self._split_doc(model)

        if ('chunk_outlines' in self.cache) and (len(self.cache['chunk_outlines']) == len(chunks)) and not force:
            return self.cache['chunk_outlines']

        pbar = tqdm(total=len(chunks), desc=f"{self.doctype.title()} chunk summarizing progress", position=1, leave=False)
        all_pts = []
        for i, chunk in enumerate(chunks):
            if ('chunk_outlines' in self.cache) and (i < len(self.cache['chunk_outlines'])) and not force:
                # resume previously interrupted process that already has *some* outlines
                pbar.update(1)
                continue
            pts = self.outline_chunk(chunk, protocol=protocol, debug=debug)
            all_pts.append(pts)
            # save intermediate progress
            self.cache['chunk_outlines'] = all_pts
            pbar.update(1)

        self.cache['chunk_outlines'] = all_pts

        if len(chunks) == 1:
            print("Only one chunk - saving as full description")
            self.cache['full_points'] = pts

        return self.cache['chunk_outlines']

    def outline_chunk(self, chunk, model='default', temperature=0, raw_response=False, protocol=0, debug=False):
        ''' Summarize a chunk of a document, into a summary and outline.

        The protocol is an integer referring to a specific prompt style.
            0: GPT is asked to outline each chunk.
            1: GPT is asked to rewrite the chunk at 25% of the length and point form. This may tease out longer outlines, which is useful at this stage.
            2: This asks for each paragraph to be described. It tries to tease out as much detail as possible, while compressing the token count. This may get too long.
            3: Similar to 2, this goes for length rather than summary. It asks for a half-length rewrite.
        '''
        if model == 'default':
            model = self.model
        _, chunks = self._split_doc(model)

        max_len = min(np.round(100 / len(chunks), 0).astype(int), 20)  # max length of summary is 100%/len(chunks) or 20% of the total length of the document

        protocols = [
        f'''Outline the following chunk of a {self.doctype}, point-by-point, with as much details as is needed to capture all arguments.
Material that may be important is: justification, background information, methods used, experimental results - with quants if available, discussion, important considerations, and broad impact.

Format the output in Markdown, under the heading 'OUTLINE'.

OUTLINE should be a finely detailed point-by-point outline of this chunk of the article. Exclude copyright information and paratextual material.

Here is the {self.doctype.lower()} to summarize:

{self.doctype.upper()}
-------

{chunk}
''',

f'''Summarize every paragraph in the following chunk of a {self.doctype}. Be as detailed as possible. The summary should be no more than {max_len}% of the length (i.e. trim it by {1-max_len}%). Capture all arguments.
Material that may be important is: justification, background information, methods used, experimental results - with quants if available, discussion, important considerations, and broad impact. Exclude copyright information and paratextual material.

Format the output in Markdown, under the heading 'OUTLINE'. Include paragraph number references for each point.

Here is the document to summarize:

{self.doctype.upper()}
-------

{chunk}
''',

f'''Paragraph by paragraph, outline every paragraph in this {self.doctype} in 5000 words, with the paragraph number cited.

Format the output in Markdown, under the heading 'OUTLINE'. Include paragraph number references for each point.

Here is the document to summarize:

{self.doctype.upper()}
-------

{chunk}
''',

f'''Rewrite the following document to half the length.

Format the output in Markdown, under the heading 'OUTLINE'. Include paragraph number references for each point.

Here is the document to summarize:

{self.doctype.upper()}
-------

{chunk}
'''

        ]
        msg = protocols[protocol]
        if debug:
            print(msg[:2000])

        result = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a very detailed document summarizer"},
                {"role": "user", "content": msg.strip()}
            ],
            temperature=temperature,
        )
        if raw_response:
            return result
        result_txt = result.choices[0].message.content
        _, pts = result_txt.split('OUTLINE')
        pts = pts.split('\n')
        return pts

    def full_summary(self, model='default', temperature=0, force=False, protocol=0):
        ''' Combine the chunk summaries into one full document summary'''

        if model == 'default':
            model = self.model

        if 'full_summary' in self.cache and not force:
            return "\n".join(self.cache['full_summary'])
        if force:
            print("Forcing a rewrite of the chunk summarization, where all the chunks are "
                  "combined into a full summary. To force the re-outlining of chunks, run "
                  "outline_all_chunks(force=True) first.")

        pbar = tqdm(total=2, desc="Preparing full summary", position=0, leave=False)
        pbar.set_description(f'Outlining {self.doctype} chunks')
        all_outlines = self.outline_all_chunks(protocol=protocol)
        pbar.update(1)
        pbar.set_description('Combining outline notes to a full summary')

        protocols = [
            '''The following is a detailed outline of multiple chunks of the same {self.doctype.lower()}. Combine them into one overall summary as fully as possible. Note the problem, approach, justification, methods employed, experimental results, broader impact, and other pertinent information.\n\n''',
            '''The following is a detailed outline of the same {self.doctype.lower()}. Rewrite to capture the most important points in 1000 words.\n\n'''
        ]
        msg = protocols[protocol]
        for i, chunk_outline in enumerate(all_outlines):
            msg += f"# CHUNK {i}\n{chunk_outline}\n\n"

        result = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a very detailed document summarizer"},
                {"role": "user", "content": msg}
            ],
            temperature=temperature,
        )
        self.cache['full_summary'] = result.choices[0].message.content.split('\n')
        pbar.update(1)
        return "\n".join(self.cache['full_summary'])

    def _parse_edit_suggestion(self, edit, html=True):
        ''' Parse a raw response line from LLM into a diff'''
        # extract first and second quote from string
        parts = edit.split('||||')
        if len(parts) != 2:
            return None
        first, second = parts
        first = first.replace('\n', ' ')
        second = second.replace('\n', ' ')
        if first.strip() == second.strip():
            return None
        diff = word_diff(first, second, html=html)
        return diff

    def edit_suggestions(self, model='default', temperature=0.3, force=False,
                         as_html=True, display_in_progress=True,
                         small_chunks=False, prompt=None):
        '''Provide edit suggestions for the {self.doctype}. This uses direct_question (which is a bit more expensive) on each chunk
        of the {self.doctype}. It then displays the diff of the original and the suggested edits, formatted as HTML. Prompt currently
        is opinionated and not customizeable, but you can use the direct_question method directly if you want to customize it.

        temperature: the temperature to use for the direct_question method. NUnlike most methods in this class, the default is not zero [default: 0.7]
        as_html: if True it returns an IPython HTML object with the diff of the original and the suggested edits - these
        can be displayed in a notebook. [default: 0.3]

        display_in_progress: if True, display the diff of each chunk as it is being processed. This is useful for reviewing the suggestions
        while the next chunk is being processed. The results are still returned at the end, so if you're in a notebook, assign to a variable
        to keep them from printing again! [default: True]

        prompt: The base request. By default, it is: "Write a numbered list of line-item edit suggestions for typo corrections, confusing sentences, poor grammar, etc."
        '''
        if ('edit_suggestions' in self.cache) and not force:
            all_edit_lines = "\n".join(self.cache['edit_suggestions'])
            all_edit_lines = re.sub(r'\n\s+(\|\|\|\|)', r'\1', all_edit_lines)  # shouldn't be needed if saved properly
            all_suggestions = [self._parse_edit_suggestion(x, html=as_html) for x in all_edit_lines.split('\n')]
            all_suggestions = [x for x in all_suggestions if x is not None]
            if as_html:
                return HTML("<br/><br/>".join(all_suggestions))
            else:
                return "\n".join(all_suggestions)

        if model == 'default':
            model = self.model
        chunksize = 800 if small_chunks else 'auto'
        _, chunks = self._split_doc(model, chunksize=chunksize)

        default_prompt = 'Write a numbered list of line-item edit suggestions for typo corrections, confusing sentences, poor grammar, etc.'
        format_instructions = 'Identity as many as you can. Respond in the format `{{original snippet}}||||{{corrected snippet}}`. Ignore smart quotes, and don\'t number the lines.'
        if not prompt:
            prompt = default_prompt
        prompt += '\n\n' + format_instructions

        all_suggestions = []
        all_edit_suggestions = ""

        all_diffs = []
        for chunk in tqdm(chunks, desc="Generating edit suggestions", position=0, leave=False):
            edit_suggestions = self.direct_question(prompt, force=force,
                                                    temperature=temperature,
                                                    md=False,
                                                    target_text=chunk,
                                                    target_chunk=None,
                                                    no_cache=True,
                                                    model=model)
            edit_suggestions = re.sub(r"\n\s*(\|\|\|\|)", r"\1", edit_suggestions)
            edit_suggestions = re.sub(r"^\d+\. ", "", edit_suggestions, flags=re.MULTILINE)
            edit_suggestions = re.sub(r"^.*?\{\{(.*?)\}\}(\|\|\|\|)\{\{(.*?)\}\}.*$", r"\1\2\3", 
                                      edit_suggestions, flags=re.MULTILINE)

            all_edit_suggestions += '\n' + edit_suggestions

            for edit in edit_suggestions.split('\n'):
                diff = self._parse_edit_suggestion(edit, html=as_html)
                if diff is None:
                    continue
                all_diffs.append(diff)

                if display_in_progress:
                    if as_html:
                        display(HTML(diff))
                    else:
                        print(diff)

        self.cache['edit_suggestions'] = all_edit_suggestions.strip().split('\n')

        if as_html:
            html = "<br/>".join(all_diffs)
            return HTML(html)
        else:
            return "\n\n".join(all_diffs)

    def full_points(self, model='default', temperature=0, force=False, protocol=0):
        ''' Combine the by-chunk key points into one set'''
        if model == 'default':
            model = self.model

        if 'full_points' in self.cache and not force:
            return "\n".join(self.cache['full_points'])

        chunk_outlines = self.outline_all_chunks(protocol=protocol)

        chunk_outlines = [pt for chunk in chunk_outlines for pt in chunk]
        in_pts = "- " + "\n- ".join(chunk_outlines)

        msg = f'''The following is a details outline of key points from multiple chunks of the same {self.doctype.lower()}.

        Combine all the points as fully and completely as possible, organized sensibly, avoiding redundancy but including as much non-redundant detail as possible. Present the most important information first. Return the information as bullet points.

        Include information such as problem and approach, justification, methods, experimental results - including quants when relevant, significance and impact, and future ramifications. Don't include points unrelated to the research, such as copyright statemements and calls for reviewers.

        Format the results as Markdown bullet points.

        # INPUT OUTLINE

        {in_pts}
        '''

        result = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a very detailed document summarizer."},
                {"role": "user", "content": msg}
            ],
            temperature=temperature
        )
        result_txt = result.choices[0].message.content
        keypts_txt = re.sub('^#.*', '', result_txt).strip()
        self.cache['full_points'] = keypts_txt.split('\n')
        return "\n".join(self.cache['full_points'])


class PDFPaperSummarizer(PaperSummarizer):
    def __init__(self, pdf_path, cache_location='match_file', *args, **kwargs):
        self.pdf_path = pdf_path

        text = self.extract_text_from_pdf()

        if cache_location == 'match_file':
            cache_location = f"{pdf_path.rsplit('.', 1)[0]}.yaml"
        elif cache_location == 'hash':
            # saves to the same directory, but append the hash of the text to the name
            p = Path(pdf_path)
            md5_hasher = hashlib.md5()
            md5_hasher.update(self.text.encode('utf-8'))
            hash = md5_hasher.hexdigest()[:6]
            cache_location = p.parent / f"{p.stem}_{hash}.yaml"

        super().__init__(text, cache_location=cache_location, *args, **kwargs)

    def extract_text_from_pdf(self):
        with open(self.pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)

            extracted_text = ""
            for page_num, page in enumerate(reader.pages):
                extracted_text += page.extract_text()

        return extracted_text


class MDPaperSummarizer(PaperSummarizer):
    def __init__(self, doc_path, cache_location='match_file', *args, **kwargs):
        self.doc_path = doc_path

        text = self.extract_text_from_md()

        if cache_location == 'match_file':
            cache_location = f"{doc_path.rsplit('.', 1)[0]}.yaml"
        elif cache_location == 'hash':
            # saves to the same directory, but append the hash of the text to the name
            p = Path(doc_path)
            md5_hasher = hashlib.md5()
            md5_hasher.update(self.text.encode('utf-8'))
            hash = md5_hasher.hexdigest()[:6]
            cache_location = p.parent / f"{p.stem}_{hash}.yaml"

        super().__init__(text, cache_location=cache_location, *args, **kwargs)

    def extract_text_from_md(self):
        with open(self.doc_path, "r") as file:
            text = file.read()
        return text


class DocxPaperSummarizer(PaperSummarizer):
    def __init__(self, doc_path, cache_location='match_file', *args, **kwargs):
        self.doc_path = doc_path

        text = self.extract_text_from_docx()

        if cache_location == 'match_file':
            cache_location = f"{doc_path.rsplit('.', 1)[0]}.yaml"
        elif cache_location == 'hash':
            # saves to the same directory, but append the hash of the text to the name
            p = Path(doc_path)
            md5_hasher = hashlib.md5()
            md5_hasher.update(self.text.encode('utf-8'))
            hash = md5_hasher.hexdigest()[:6]
            cache_location = p.parent / f"{p.stem}_{hash}.yaml"

        super().__init__(text, cache_location=cache_location, *args, **kwargs)

    def extract_text_from_docx(self):
        from docx import Document
        document = Document(self.doc_path)

        text = ""
        for paragraph in document.paragraphs:
            formatted_paragraph = self.markdown_formatting(paragraph)

            # Write the formatted text to the file
            text += formatted_paragraph.strip() + "\n\n"
        return text.strip()

    def markdown_formatting(self, paragraph):
        """
        This function receives a paragraph object and iterates over its runs,
        formatting bold and italic text with Markdown and returning the resulting string.
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        bold, bolditalic, italic, regular = 0, 0, 0, 0
        text = ""
        # count how much of the text is a certain style (for making heading inferences)
        for i, run in enumerate(paragraph.runs):
            text += run.text
            if run.text.strip() == '':
                continue
            if run.bold:
                bold += len(run.text)
            elif run.bold and run.italic:
                bolditalic += len(run.text)
            elif run.italic:
                italic += len(run.text)
            else:
                regular += len(run.text)

        total_len = len(text.strip())

        if paragraph.style.name == "Normal":
            # make guesses for a 'Normal formatted' paragraph
            if total_len <= bold:
                if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    # Heading: Bold and Centered
                    return f"# {text.strip()}"
                else:
                    # Subheading: Bold
                    return f"## {text.strip()}"
            elif total_len <= bolditalic:
                # Subsubheading: Bold and Italic
                return f"### {text.strip()}"
            elif total_len <= italic:
                # Subsubsubheading: Italic
                return f"#### {text.strip()}"
            else:
                # Apply formatting to individual runs
                formatted_text = ""
                for run in paragraph.runs:
                    end_spaces = run.text[len(run.text.rstrip()):]
                    start_spaces = run.text[:len(run.text) - len(run.text.lstrip())]
                    if run.text.strip() == '':
                        formatted_text += run.text
                    elif run.bold and run.italic:
                        # Both bold and italic
                        formatted_text += f"{start_spaces}***{run.text.strip()}***{end_spaces}"
                    elif run.bold:
                        # Only bold
                        formatted_text += f"{start_spaces}**{run.text.strip()}**{end_spaces}"
                    elif run.italic:
                        # Only italic
                        formatted_text += f"{start_spaces}_{run.text.strip()}_{end_spaces}"
                    else:
                        # Regular text
                        formatted_text += run.text
                # fix any formatting within word
                formatted_text = formatted_text.replace('****', '').replace('’', "'").replace('“', '"').replace('”', '"')
                formatted_text = re.sub(r'\*\* +\*\*', r' ', formatted_text)
                # formatted_text = re.sub('(\W)([\*\_]{1,2})', r'\1 \2', formatted_text)
                # formatted_text = re.sub('(\w)[\*\_]{1,2}(\w)', r'\1\2', formatted_text)
                return formatted_text

        elif 'Heading' in paragraph.style.name:
            level = paragraph.style.name.split(' ')[1]  # Assume style is like 'Heading 1'
            return f"{'#' * int(level)} {text.strip()}"

        else:
            return text


class WebPageSummarizer(PaperSummarizer):
    def __init__(self, url, cache_location='match_file', *args, **kwargs):
        self.url = url
        self.title = None

        text = self.extract_text_from_url()

        if cache_location == 'match_file':
            cache_location = f"{url.split('://')[-1].replace('/', '_')}.yaml"
        elif cache_location == 'hash':
            # saves to the same directory, but append the hash of the text to the name
            p = Path(url)
            md5_hasher = hashlib.md5()
            md5_hasher.update(self.text.encode('utf-8'))
            hash = md5_hasher.hexdigest()[:6]
            cache_location = p.parent / f"{p.stem}_{hash}.yaml"

        super().__init__(text, cache_location=cache_location, *args, **kwargs)

    def extract_text_from_url(self):
        import requests
        from readability import Document
        response = requests.get(self.url)
        doc = Document(response.content)
        self.title = doc.title()
        html = doc.summary()
        return self.parse_html_to_markdown(html)

    def parse_html_to_markdown(self, html):
        import html2text
        return html2text.html2text(html)
